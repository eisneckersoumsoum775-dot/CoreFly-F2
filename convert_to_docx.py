"""
CoreFly F2 软件设计计划书 → Word
- 盒子框图 → Word 单列表格 (公式/代码/结构图)
- 多列表 → Word 多列表格
- 标题 → Heading 1/2/3 (TOC 可跳转)
- 希腊字母 → Unicode
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r'F:\Zhuomian\CoreFly F2\MDK-ARM\软件设计计划书.txt'
OUTPUT = r'F:\Zhuomian\CoreFly F2\CoreFly F2 软件设计计划书.docx'

FONT_BODY = '微软雅黑'; FONT_MONO = 'Consolas'; FONT_H = '微软雅黑'
C_H1=RGBColor(0x1A,0x1A,0x2E); C_H2=RGBColor(0x0D,0x3B,0x66); C_H3=RGBColor(0x2C,0x3E,0x50)
C_TEXT=RGBColor(0x22,0x22,0x22); C_GREY=RGBColor(0x99,0x99,0x99)
C_RED=RGBColor(0xE9,0x45,0x60); C_WHITE=RGBColor(0xFF,0xFF,0xFF)
C_MONO=RGBColor(0x2D,0x2D,0x2D)
SZ_H1=Pt(16);SZ_H2=Pt(13);SZ_H3=Pt(11);SZ_B=Pt(10);SZ_M=Pt(9.5);SZ_T=Pt(9)

BOX_CHARS = set('┌└┐┘│├─┬┴┼┤╔╚╗╝║═╠╦╩╬╣┊┋┃┄┅┈┉')

GREEK = {
    'alpha':'α','beta':'β','gamma':'γ','delta':'δ','epsilon':'ε','zeta':'ζ',
    'eta':'η','theta':'θ','iota':'ι','kappa':'κ','lambda':'λ','mu':'μ',
    'nu':'ν','xi':'ξ','pi':'π','rho':'ρ','sigma':'σ','tau':'τ',
    'upsilon':'υ','phi':'φ','chi':'χ','psi':'ψ','omega':'ω',
    'Alpha':'Α','Beta':'Β','Gamma':'Γ','Delta':'Δ','Theta':'Θ',
    'Sigma':'Σ','Omega':'Ω','Pi':'Π','Phi':'Φ',
}
def _greek(t):
    for k,v in sorted(GREEK.items(),key=lambda x:-len(x[0])): t=t.replace(k,v)
    return t

# ============================================================
def _font(run, name, size, bold=False, color=None, super_=False, sub_=False):
    run.font.size=size;run.font.bold=bold;run.font.name=name
    if color:run.font.color.rgb=color
    if super_:run.font.superscript=True
    if sub_:run.font.subscript=True
    rPr=run._element.get_or_add_rPr()
    rF=rPr.find(qn('w:rFonts'))
    if rF is None:rF=OxmlElement('w:rFonts');rPr.insert(0,rF)
    rF.set(qn('w:eastAsia'),name);rF.set(qn('w:ascii'),name);rF.set(qn('w:hAnsi'),name)

def _run(p,text,name=FONT_BODY,size=SZ_B,bold=False,color=None):
    r=p.add_run(text);_font(r,name,size,bold,color);return r

def _sp(p,b=0,a=0,lh=1.15):
    pf=p.paragraph_format;pf.space_before=Pt(b);pf.space_after=Pt(a);pf.line_spacing=lh

def _hdr(doc,text,lv):
    p=doc.add_paragraph()
    c={1:C_H1,2:C_H2,3:C_H3}[lv];s={1:SZ_H1,2:SZ_H2,3:SZ_H3}[lv]
    bs={1:14,2:10,3:6}[lv];af={1:6,2:4,3:2}[lv]
    _sp(p,b=bs,a=af);p.style=doc.styles[f'Heading {lv}']
    _run(p,text,name=FONT_H,size=s,bold=True,color=c)

def _body(doc,text,indent=0):
    p=doc.add_paragraph();_sp(p,b=0,a=1)
    if indent>=1:p.paragraph_format.left_indent=Cm(0.7)
    if indent>=2:p.paragraph_format.left_indent=Cm(1.4)
    _run(p,_greek(text),size=SZ_B,color=C_TEXT)

def _hr(doc):
    p=doc.add_paragraph();_sp(p,b=4,a=4)
    pBdr=OxmlElement('w:pBdr');bt=OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single');bt.set(qn('w:sz'),'4');bt.set(qn('w:space'),'1');bt.set(qn('w:color'),'CCCCCC')
    pBdr.append(bt);p._element.get_or_add_pPr().append(pBdr)

def _gap(doc,pt=4):
    p=doc.add_paragraph();_sp(p);_run(p,' ',size=Pt(pt))

def _cell_shade(cell,hex_color):
    tcPr=cell._element.get_or_add_tcPr();sd=OxmlElement('w:shd')
    sd.set(qn('w:val'),'clear');sd.set(qn('w:color'),'auto');sd.set(qn('w:fill'),hex_color)
    tcPr.append(sd)

# ============================================================
def _is_box(s):s=s.strip();return len(s)>5 and all(c in BOX_CHARS or c==' ' for c in s)
def _is_top(s):s=s.strip();return len(s)>5 and s[0] in '┌╔' and _is_box(s)
def _is_bot(s):s=s.strip();return len(s)>5 and s[0] in '└╚' and _is_box(s)
def _is_tbltop(s):return _is_top(s) and ('┬' in s or '╦' in s or '╤' in s)
def _is_midsep(s):s=s.strip();return len(s)>5 and ('├' in s or '┼' in s) and _is_box(s)
def _is_tblrow(s):s=s.strip();return s.startswith('│') and s.endswith('│') and not _is_midsep(s)

def _extract_box_block(lines, start):
    inner=[];i=start+1;total=len(lines)
    while i<total:
        s=lines[i].strip()
        if _is_bot(s):return inner,i+1
        if s.startswith('│'):inner.append(s)
        else:break
        i+=1
    return inner,i

def _is_formula_line(text):
    """判断一行是否为公式（有下标/上标/等号/数学符号）"""
    s=text.strip()
    if any(k in s for k in ['=','_','^','×','√','→','∂','d/dt','(','/']):return True
    return False

def _add_formula(p, text):
    """解析公式文本，生成带上下标的 Word runs"""
    text=_greek(text)  # alpha→α, beta→β ...
    i=0;n=len(text);buf=[]

    def _flush(buf, sub=False, sup=False):
        if not buf:return
        s=''.join(buf);buf.clear()
        r=p.add_run(s)
        sz=Pt(SZ_M.pt*0.85) if (sub or sup) else SZ_M
        _font(r,FONT_MONO,sz,color=C_MONO,sub_=sub,super_=sup)

    def _is_ident(c):
        return c.isalnum() or c in 'αβγδεζηθικλμνξπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΠΡΣΤΥΦΧΨΩ'

    while i<n:
        c=text[i]

        # 下标: X_YYY
        if c=='_' and buf and i+1<n and not text[i-1].isspace():
            # 收集下标内容
            i+=1;j=i
            while j<n and (_is_ident(text[j]) or text[j] in '.-'):
                j+=1
            sub_text=text[i:j]
            if sub_text:
                _flush(buf)
                r=p.add_run(_greek(sub_text))
                _font(r,FONT_MONO,Pt(SZ_M.pt*0.82),color=C_MONO,sub_=True)
                i=j;continue
            else:
                buf.append('_');i+=1;continue

        # 上标: X^Y
        if c=='^' and buf and i+1<n:
            i+=1;j=i
            while j<n and (_is_ident(text[j]) or text[j] in '.-'):j+=1
            sup_text=text[i:j]
            if sup_text:
                _flush(buf)
                r=p.add_run(sup_text)
                _font(r,FONT_MONO,Pt(SZ_M.pt*0.82),color=C_MONO,super_=True)
                i=j;continue
            else:
                buf.append('^');i+=1;continue

        buf.append(c);i+=1

    _flush(buf)

def _strip_box_edge(s):
    s=s.strip()
    if s.startswith('│'):s=s[1:]
    if s.endswith('│'):s=s[:-1]
    return s

def _box_to_table(doc, inner_lines):
    """盒子内容行 → Word 1列表格"""
    if not inner_lines:return
    content=[]
    for line in inner_lines:
        s=_strip_box_edge(line)
        if s and not all(c in '─═━' for c in s.strip()):
            content.append(s)
    if not content:return

    has_code=any(s.strip().startswith(('//','/*','*/','void ','float ','int ',
        'static ','typedef','struct ','if (','for (','while (','return ',
        '    ','#include','#define','}','{','// =',' * @')) for s in content)

    bg='F0F0F0' if has_code else 'F5F5F5'

    table=doc.add_table(rows=len(content),cols=1);table.autofit=True
    tbl=table._tbl;tblPr=tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW=OxmlElement('w:tblW');tblW.set(qn('w:w'),'5000');tblW.set(qn('w:type'),'pct')
    tblPr.append(tblW)
    tblBorders=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        border=OxmlElement(f'w:{edge}')
        border.set(qn('w:val'),'single');border.set(qn('w:sz'),'4');border.set(qn('w:space'),'0');border.set(qn('w:color'),'BBBBBB')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for i,txt in enumerate(content):
        cell=table.cell(i,0);cell.text=''
        p=cell.paragraphs[0];_sp(p,b=0,a=0,lh=1.15)
        _cell_shade(cell,bg)
        if has_code:
            _run(p,_greek(txt),name=FONT_MONO,size=SZ_M,color=C_MONO)
        else:
            _add_formula(p,txt)
    _gap(doc,4)

# ============================================================
def _extract_table(lines, start):
    i=start;total=len(lines)
    if i>=total:return None
    if not _is_tbltop(lines[i].strip()):return None
    i+=1;hdr=[];dat=[];past=False
    while i<total:
        s=lines[i].strip()
        if _is_bot(s):i+=1;return (hdr,dat,i) if past else (hdr,[],i)
        if _is_midsep(s):past=True;i+=1;continue
        if _is_tblrow(s):
            if past:dat.append(s)
            else:hdr.append(s)
            i+=1;continue
        if _is_top(s) and i>start+1:return (hdr,dat,i) if past else (hdr,[],i)
        break
    return (hdr,dat,i)

def _render_multi_table(doc,hdr,dat):
    def _pr(s):
        s=s.strip()
        if s.startswith('│'):s=s[1:]
        if s.endswith('│'):s=s[:-1]
        return [c.strip() for c in s.split('│')]
    parsed=[_pr(r) for r in (hdr+dat)]
    mc=max(len(r) for r in parsed)if parsed else 1
    for r in parsed:
        while len(r)<mc:r.append('')
    nr=(1 if hdr else 0)+len(dat)
    if nr==0:return
    tbl=doc.add_table(rows=nr,cols=mc,style='Table Grid');tbl.autofit=True
    if hdr:
        hp=_pr(hdr[0])
        while len(hp)<mc:hp.append('')
        for j,txt in enumerate(hp):
            c=tbl.cell(0,j);c.text=''
            _run(c.paragraphs[0],txt,name=FONT_H,size=SZ_T,bold=True,color=C_WHITE)
            _cell_shade(c,'1A1A2E')
    off=1 if hdr else 0
    for i,row in enumerate(dat):
        rp=_pr(row)
        while len(rp)<mc:rp.append('')
        for j,txt in enumerate(rp):
            c=tbl.cell(off+i,j);c.text=''
            _run(c.paragraphs[0],txt,size=SZ_T,color=C_TEXT)
    _gap(doc,4)

# ============================================================
def _detect(s):
    s=s.strip()
    if not s:return None
    if s.startswith('===') and ('部分' in s or '附录' in s):return (1,s.strip('= ').strip())
    m=re.match(r'^(第[一二三四五六七八九十百]+部分[：:]\s*.+)',s)
    if m:return(1,m.group(1))
    m=re.match(r'^(第\d+阶段[：:]\s*.+)',s)
    if m:return(1,m.group(1))
    m=re.match(r'^(附录[A-Z][：:]\s*.+)',s)
    if m:return(1,m.group(1))
    m=re.match(r'^([一二三四五六七八九十]、\s*.+)',s)
    if m and len(s)<100:return(2,m.group(1))
    m=re.match(r'^(\d+\.\d+\.\d+)\s{1,3}(.+)',s)
    if m and len(s)<100:return(3,f'{m.group(1)}  {m.group(2)}')
    m=re.match(r'^(\d+\.\d+)\s{1,3}(.+)',s)
    if m and len(s)<100:return(2,f'{m.group(1)}  {m.group(2)}')
    m=re.match(r'^\s{2,4}(任务\d[\d\.]*\s+.+)',s)
    if m:return(3,m.group(1).strip())
    return None

_CODE=('// ','/* ','*/ ','#include ','#define ','void ','float ','int ','uint',
    'static ','typedef ','struct ','if (','for (','while (','return ','else ',
    'case ','break;','    //','    float','    int','    uint',
    '    if','    for','    return','    }','    err_','    z_','    E_',
    '// =',' * @','* @','HAL_','    motor','m->','motor->','smo->','bmi_')
def _is_code(s):
    s=s.strip()
    for kw in _CODE:
        if s.startswith(kw):return True
    return False

# ============================================================
def build():
    with open(INPUT,'r',encoding='utf-8') as f:
        lines=[l.rstrip() for l in f.readlines()]

    doc=Document()
    st=doc.styles['Normal'];st.font.name=FONT_BODY;st.font.size=SZ_B
    st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT_BODY)
    for lv in[1,2,3]:
        hs=doc.styles[f'Heading {lv}'];hs.font.name=FONT_H;hs.font.bold=True
        hs.font.size={1:SZ_H1,2:SZ_H2,3:SZ_H3}[lv]
        hs.font.color.rgb={1:C_H1,2:C_H2,3:C_H3}[lv]
        hs.element.rPr.rFonts.set(qn('w:eastAsia'),FONT_H)
        hs.paragraph_format.space_before=Pt({1:14,2:10,3:6}[lv])
        hs.paragraph_format.space_after=Pt({1:6,2:4,3:2}[lv])

    sec=doc.sections[0]
    sec.page_width=Cm(21.0);sec.page_height=Cm(29.7)
    sec.left_margin=Cm(2.2);sec.right_margin=Cm(1.8)
    sec.top_margin=Cm(1.8);sec.bottom_margin=Cm(1.8)

    for _ in range(6):_gap(doc,12)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(p,'CoreFly F2',name=FONT_H,size=Pt(32),bold=True,color=C_H1)
    _gap(doc,8)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(p,'飞控软件设计计划书',name=FONT_H,size=Pt(22),bold=True,color=C_H2)
    _gap(doc,6)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(p,'━'*30,size=Pt(10),color=C_RED)
    _gap(doc,10)
    for info in ['版本：v2.0 (详尽版)  |  日期：2026-05-23',
                 '适用硬件：STM32F407VET6 + FD6288Q + AON7418 × 24',
                 '适用电机：4路 BLDC (无感 FOC 驱动)',
                 '','—— 逐行对着写就能实现的开发手册 ——']:
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p,info,size=Pt(10),color=C_GREY)
    doc.add_page_break()

    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(p,'目    录',name=FONT_H,size=Pt(20),bold=True,color=C_H1)
    _gap(doc,8)
    p=doc.add_paragraph()
    r=p.add_run();fb=OxmlElement('w:fldChar');fb.set(qn('w:fldCharType'),'begin');r._element.append(fb)
    r2=p.add_run();it=OxmlElement('w:instrText');it.set(qn('xml:space'),'preserve')
    it.text=' TOC \\o "1-3" \\h \\z \\u ';r2._element.append(it)
    r3=p.add_run();fs=OxmlElement('w:fldChar');fs.set(qn('w:fldCharType'),'separate');r3._element.append(fs)
    r4=p.add_run();r4.text='（打开文档后，右键此处 → 更新域 → 更新整个目录，即可生成可跳转的目录）'
    _font(r4,FONT_BODY,Pt(9),color=C_GREY)
    r5=p.add_run();fe=OxmlElement('w:fldChar');fe.set(qn('w:fldCharType'),'end');r5._element.append(fe)
    doc.add_page_break()

    i=0;total=len(lines)
    while i<total and (lines[i].strip().startswith('===') or lines[i].strip()==''):i+=1

    while i<total:
        line=lines[i];s=line.strip()
        if not s:i+=1;continue
        if '目' in s and '录' in s and i<60:
            i+=1
            while i<total and not lines[i].strip().startswith('==='):i+=1
            continue
        h=_detect(s)
        if h:
            lv,title=h
            if lv==1:_hr(doc);_hdr(doc,title,1);_hr(doc)
            else:_hdr(doc,title,lv)
            i+=1;continue
        if s.startswith('===') and i>20:_hr(doc);i+=1;continue
        if(s.startswith('---')and len(s)>10)or(s.startswith('───')and len(s)>4):_hr(doc);i+=1;continue
        if _is_tbltop(s):
            tbl=_extract_table(lines,i)
            if tbl:
                hdr,dat,ei=tbl
                if dat:_render_multi_table(doc,hdr,dat)
                else:
                    inner,ei2=_extract_box_block(lines,i);_box_to_table(doc,inner);ei=ei2
                i=ei;continue
        if _is_top(s):
            inner,ei=_extract_box_block(lines,i)
            _box_to_table(doc,inner)
            i=ei;continue
        if _is_code(s):
            cl=[s];i+=1
            while i<total:
                ns=lines[i].strip()
                if ns and (_is_code(ns)or(ns and ns[0].isspace()and not ns.startswith('│'))):
                    cl.append(lines[i].rstrip());i+=1
                else:break
            _box_to_table(doc,cl);continue
        if s.startswith(('目标:','前提:','验证标准:','⚠','说明：')):
            _body(doc,s);doc.paragraphs[-1].runs[0].bold=True;i+=1;continue
        if s.startswith('□'):_body(doc,'☐ '+s[1:].strip(),indent=1);i+=1;continue
        indent=0
        if line.startswith('    ')or line.startswith('\t'):
            indent=1
            if line.startswith('        '):indent=2
        _body(doc,s,indent=indent);i+=1

    doc.save(OUTPUT)
    h1=h2=h3=0
    for p in doc.paragraphs:
        if p.style.name=='Heading 1':h1+=1
        elif p.style.name=='Heading 2':h2+=1
        elif p.style.name=='Heading 3':h3+=1
    print(f'OK: {os.path.basename(OUTPUT)}')
    print(f'Size: {os.path.getsize(OUTPUT)/1024:.0f} KB')
    print(f'Headings: H1x{h1} H2x{h2} H3x{h3} ({h1+h2+h3})')
    print(f'Tables: {len(doc.tables)}')

if __name__=='__main__':
    build()
